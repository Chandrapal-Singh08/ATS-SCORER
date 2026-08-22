import io
from pathlib import Path
from typing import Optional, Tuple

import PyPDF2
import pdfplumber
from docx import Document

from backend.core.config import (
    MAX_FILE_SIZE_BYTES,
    MAX_FILE_SIZE_MB,
)
from backend.utils.file_utils import (
    FileParsingError,
    TextExtractionError,
    FileUploadError,
    log_error,
    log_warning,
    log_info,
    with_fallback,
)


class FileParsingError(Exception):
    pass


class FileValidationError(Exception):
    pass


# -------------------------------------------------------
# FILE VALIDATION (Render Compatible - No python-magic)
# -------------------------------------------------------
def validate_file(
    file_data: bytes,
    filename: str,
) -> Tuple[bool, str, Optional[str]]:

    file_size_bytes = len(file_data)

    if file_size_bytes > MAX_FILE_SIZE_BYTES:
        size_mb = file_size_bytes / (1024 * 1024)
        return (
            False,
            f"File size ({size_mb:.2f} MB) exceeds the maximum of "
            f"{MAX_FILE_SIZE_MB} MB. Please upload a smaller file.",
            None,
        )

    if file_size_bytes == 0:
        return False, "Uploaded file is empty.", None

    extension = Path(filename).suffix.lower()

    if extension == ".pdf":
        file_type = "pdf"

    elif extension == ".docx":
        file_type = "docx"

    elif extension == ".doc":
        file_type = "doc"

    else:
        return (
            False,
            "Unsupported file type. Please upload a PDF or DOCX resume.",
            None,
        )

    return True, "", file_type


# -------------------------------------------------------
# PDF HYPERLINK EXTRACTION
# -------------------------------------------------------
def _extract_pdf_hyperlinks(file_data: bytes) -> str:
    urls = []

    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_data))

        for page in reader.pages:

            if "/Annots" not in page:
                continue

            for annot_ref in page["/Annots"]:

                try:
                    annot = annot_ref.get_object()

                    if annot.get("/Subtype") != "/Link":
                        continue

                    action = annot.get("/A", {})
                    uri = action.get("/URI", "")

                    if uri:

                        if isinstance(uri, bytes):
                            uri = uri.decode("utf-8", errors="ignore")

                        uri = uri.strip()

                        if uri.startswith("http"):
                            urls.append(uri)

                except Exception:
                    pass

    except Exception:
        pass

    return "\n".join(urls)


# -------------------------------------------------------
# PDF EXTRACTION USING PDFPLUMBER
# -------------------------------------------------------
def _extract_pdf_with_pdfplumber(file_data: bytes) -> str:

    text = ""

    with pdfplumber.open(io.BytesIO(file_data)) as pdf:

        for page in pdf.pages:
            page_text = page.extract_text()

            if page_text:
                text += page_text + "\n"

    if not text.strip():
        raise TextExtractionError(
            "pdfplumber extracted no text",
            user_message="No text could be extracted from this PDF.",
        )

    hyperlinks = _extract_pdf_hyperlinks(file_data)

    if hyperlinks:
        text = text.strip() + "\n" + hyperlinks

    return text.strip()


# -------------------------------------------------------
# PDF EXTRACTION USING PYPDF2 (Fallback)
# -------------------------------------------------------
def _extract_pdf_with_pypdf2(file_data: bytes) -> str:

    text = ""

    reader = PyPDF2.PdfReader(io.BytesIO(file_data))

    for page in reader.pages:
        page_text = page.extract_text()

        if page_text:
            text += page_text + "\n"

    if not text.strip():
        raise TextExtractionError(
            "PyPDF2 extracted no text",
            user_message="No text could be extracted from this PDF.",
        )

    hyperlinks = _extract_pdf_hyperlinks(file_data)

    if hyperlinks:
        text = text.strip() + "\n" + hyperlinks

    return text.strip()


# -------------------------------------------------------
# PUBLIC PDF EXTRACTION FUNCTION
# -------------------------------------------------------
def extract_text_from_pdf(file_data: bytes) -> str:

    try:
        result, used_fallback = with_fallback(
            _extract_pdf_with_pdfplumber,
            _extract_pdf_with_pypdf2,
            file_data,
            log_fallback=True,
        )

        if used_fallback:
            log_info(
                "PDF extraction succeeded using PyPDF2 fallback.",
                context="resume_parser",
            )

        return result

    except Exception as e:
        log_error(e, context="extract_text_from_pdf")

        raise FileParsingError(
            "Failed to extract text from PDF. The PDF may be corrupted, "
            "password protected, or contain only scanned images."
        ) from e


# -------------------------------------------------------
# DOCX EXTRACTION
# -------------------------------------------------------
def extract_text_from_docx(file_data: bytes) -> str:

    try:
        doc = Document(io.BytesIO(file_data))

        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        text = "\n".join(text_parts)

        if not text.strip():
            raise FileParsingError(
                "No text could be extracted from the DOCX file."
            )

        try:
            for rel in doc.part.rels.values():

                if "hyperlink" in rel.reltype.lower():

                    url = rel._target

                    if isinstance(url, str) and url.startswith("http"):
                        text += "\n" + url

        except Exception:
            pass

        log_info(
            f"Extracted {len(text)} chars from DOCX",
            context="resume_parser",
        )

        return text.strip()

    except FileParsingError:
        raise

    except Exception as e:
        log_error(e, context="extract_text_from_docx")

        raise FileParsingError(
            "Failed to extract text from DOCX. The document may be corrupted."
        ) from e


# -------------------------------------------------------
# LEGACY DOC
# -------------------------------------------------------
def extract_text_from_doc(file_data: bytes) -> str:

    raise FileParsingError(
        "Legacy .doc format is not supported. "
        "Please convert it to .docx or .pdf."
    )


# -------------------------------------------------------
# DISPATCH EXTRACTION
# -------------------------------------------------------
def extract_text(file_data: bytes, file_type: str) -> str:

    if file_type == "pdf":
        return extract_text_from_pdf(file_data)

    if file_type == "docx":
        return extract_text_from_docx(file_data)

    if file_type == "doc":
        return extract_text_from_doc(file_data)

    raise FileValidationError(f"Invalid file type: {file_type}")


# -------------------------------------------------------
# MAIN PARSER
# -------------------------------------------------------
def parse_resume_file(
    file_data: bytes,
    filename: str,
) -> Tuple[str, dict]:

    log_info(
        f"Parsing resume: {filename}",
        context="parse_resume_file",
    )

    # Validate file
    is_valid, error_msg, file_type = validate_file(file_data, filename)

    if not is_valid:
        log_warning(error_msg, context="parse_resume_file")
        raise FileValidationError(error_msg)

    # Extract text
    try:
        text = extract_text(file_data, file_type)

        log_info(
            f"Extracted {len(text)} characters from {filename}",
            context="parse_resume_file",
        )

    except FileParsingError:
        raise

    except Exception as e:
        log_error(e, context="parse_resume_file")

        raise FileParsingError(
            "Unexpected error while processing resume."
        ) from e

    metadata = {
        "filename": filename,
        "file_type": file_type,
        "file_size_bytes": len(file_data),
        "text_length": len(text),
        "success": True,
    }

    return text, metadata